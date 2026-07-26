import tempfile
from pathlib import Path
from docx import Document as DocxDocument
from fpdf import FPDF
from src.loaders import _detect_format, _read_normalized, load_file


def test_detect_format_md():
    assert _detect_format(Path("doc.md")) == "md"

def test_detect_format_txt():
    assert _detect_format(Path("notes.txt")) == "txt"

def test_detect_format_html():
    assert _detect_format(Path("page.html")) == "html"

def test_detect_format_docx():
    assert _detect_format(Path("notes.docx")) == "docx"

def test_detect_format_pdf():
    assert _detect_format(Path("report.pdf")) == "pdf"


def test_read_normalized_strips_crlf(): #clean the file
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
        tmp.write(b"line one\r\nline two\r\n")
    f = Path(tmp.name)
    result = _read_normalized(f)
    assert result == "line one\nline two\n"
    f.unlink()


def test_load_markdown_preserves_headings(tmp_path):
    f = tmp_path / "doc.md"
    f.write_text("# Title\n\nSome text with ERR_2043.")
    doc = load_file(f, base_dir=tmp_path)
    assert doc.fmt == "md"
    assert doc.source_name == "doc.md"
    assert doc.doc_id == "doc.md"
    assert "# Title" in doc.clean_text
    assert "ERR_2043" in doc.clean_text


def test_load_txt_passthrough(tmp_path):
    f = tmp_path / "notes.txt"
    f.write_text("plain text, no markup")
    doc = load_file(f, base_dir=tmp_path)
    assert doc.fmt == "txt"
    assert "plain text" in doc.clean_text


def test_load_html_converts_headings_and_strips_script(tmp_path):
    f = tmp_path / "doc.html"
    f.write_text(
        "<html><body><h1>Title</h1><script>evil()</script><p>Hi there</p></body></html>"
    )
    doc = load_file(f, base_dir=tmp_path)
    assert doc.fmt == "html"
    assert "# Title" in doc.clean_text
    assert "Hi there" in doc.clean_text
    assert "evil" not in doc.clean_text


def test_load_html_converts_nested_heading_levels(tmp_path):
    f = tmp_path / "doc2.html"
    f.write_text("<html><body><h2>Section</h2><p>Body text</p></body></html>")
    doc = load_file(f, base_dir=tmp_path)
    assert "## Section" in doc.clean_text


def test_load_docx_converts_headings_and_extracts_paragraphs(tmp_path):
    f = tmp_path / "doc.docx"
    docx_doc = DocxDocument()
    docx_doc.add_heading("Title", level=1)
    docx_doc.add_paragraph("Body text with ERR_2043.")
    docx_doc.add_heading("Subsection", level=2)
    docx_doc.save(f)

    doc = load_file(f, base_dir=tmp_path)
    assert doc.fmt == "docx"
    assert "# Title" in doc.clean_text
    assert "## Subsection" in doc.clean_text
    assert "Body text with ERR_2043." in doc.clean_text
    assert "Title" in doc.raw_text
    assert "# Title" not in doc.raw_text