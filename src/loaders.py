import re
from pathlib import Path

from bs4 import BeautifulSoup

from src.models import Document

_DOCX_HEADING_PATTERN = re.compile(r"Heading (\d+)")


def _detect_format(path: Path) -> str:
    p = path
    return p.suffix[1:]


def _read_normalized(path: Path) -> str:
    """Read a text file and normalize Windows line endings to \n."""
    raw = path.read_text()
    return raw.replace("\r\n", "\n")


def _html_to_clean_text(html: str) -> str:
    """Strip script/style, convert h1-h6 to markdown headings, extract text."""
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style"]):
        tag.decompose()

    for level in range(1, 7):
        for heading in soup.find_all(f"h{level}"):
            heading.replace_with(f"\n{'#' * level} {heading.get_text(strip=True)}\n")

    return soup.get_text(separator=" ").strip()


def _docx_to_texts(path: Path) -> tuple[str, str]:
    """Extract paragraphs from a .docx file. Word's built-in "Heading N"
    paragraph styles become markdown '#' * N headings so the chunkers'
    heading-based structure detection works the same as for md/html.
    Returns (raw_text, clean_text): raw_text is the plain paragraph text
    with no heading markup; clean_text has headings converted.
    """
    from docx import Document as DocxDocument

    docx_doc = DocxDocument(str(path))

    raw_lines = []
    clean_lines = []
    for paragraph in docx_doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        raw_lines.append(text)

        style_name = paragraph.style.name if paragraph.style else ""
        match = _DOCX_HEADING_PATTERN.match(style_name)
        if match:
            level = min(int(match.group(1)), 6)
            clean_lines.append(f"{'#' * level} {text}")
        else:
            clean_lines.append(text)

    return "\n\n".join(raw_lines), "\n\n".join(clean_lines)


def load_file(path: Path, base_dir: Path) -> Document:
    """Load a single file and return a normalized Document."""
    fmt = _detect_format(path)
    source_name = path.relative_to(base_dir).as_posix()

    if fmt in ("md", "txt"):
        clean_text = _read_normalized(path)
        raw_text = path.read_text()
        metadata = {}
    elif fmt == "html":
        raw_text = path.read_text()
        clean_text = _html_to_clean_text(raw_text)
        metadata = {}
    elif fmt == "docx":
        raw_text, clean_text = _docx_to_texts(path)
        metadata = {}
    else:
        raise ValueError(f"unsupported format: {fmt}")

    return Document(
        doc_id=source_name,
        source_path=str(path),
        source_name=source_name,
        fmt=fmt,
        raw_text=raw_text,
        clean_text=clean_text,
        metadata=metadata,
    )
